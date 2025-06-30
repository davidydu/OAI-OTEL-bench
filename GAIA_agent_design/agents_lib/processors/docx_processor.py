from __future__ import annotations

from typing import List

from ..file_router import Attachment

from docx import Document


class DocxProcessorAgent:
    """Extract text from .docx files."""

    def process(self, att: Attachment) -> str:
        """Extract text from a DOCX file."""
        if att.path.exists():
            doc = Document(att.path)
        else:  # in-memory file (e.g. from a ZIP archive)
            from io import BytesIO

            doc = Document(BytesIO(att.bytes))
        texts: List[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    texts.append(row_text)
        return "\n".join(texts)
